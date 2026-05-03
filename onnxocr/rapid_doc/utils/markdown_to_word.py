# -*- coding: utf-8 -*-
"""
Markdown 转换工具模块
支持 Markdown 转 docx

技术栈:
- Markdown 转 docx: pypandoc (基于 Pandoc)
"""
import os
from pathlib import Path
from typing import Optional

from loguru import logger

# ===================== 依赖检测 =====================
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

# python-docx 用于后处理 Word 文档（添加表格边框等）
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

def _add_table_borders(docx_path: str) -> None:
    """
    为 Word 文档中的所有表格添加边框
    使用 python-docx 操作 OOXML 直接设置表格边框属性
    """
    doc = Document(docx_path)
    
    for table in doc.tables:
        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # 创建表格边框元素
        tblBorders = OxmlElement('w:tblBorders')
        
        # 定义边框类型
        border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        
        for border_name in border_types:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # 单线边框
            border.set(qn('w:sz'), '4')        # 边框宽度 0.5pt (4 = 0.5pt * 8)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 黑色
            tblBorders.append(border)
        
        # 移除已有的边框设置（如果有）
        for existing in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(existing)
        
        tblPr.append(tblBorders)
    
    # 保存文档
    doc.save(docx_path)


def _set_fonts(docx_path: str, chinese_font: str = '宋体', latin_font: str = 'Times New Roman') -> None:
    """
    设置 Word 文档的默认字体：中文使用宋体，英文/数字使用 Times New Roman
    """
    if not PYTHON_DOCX_AVAILABLE:
        return

    doc = Document(docx_path)

    # 1. 修改正文默认样式 (Normal Style)
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)  # 可选：设置默认字号

    # 设置西文字体
    font.name = latin_font
    # 设置中文字体 (必须通过 oxml 直接操作)
    rFonts = font._element.rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), chinese_font)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)

    # 2. 遍历所有段落和运行内容，确保字体应用
    # (Pandoc 生成的内容有时会带有具体的格式覆盖，直接改 Style 可能不彻底)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = latin_font
            run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), chinese_font)

    # 3. 处理表格中的字体
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = latin_font
                        run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), chinese_font)

    doc.save(docx_path)


from docx.shared import RGBColor  # 需要导入这个类


def _fix_styles(docx_path: str):
    """
    统一修改文档样式：将所有标题颜色改为黑色
    """
    if not PYTHON_DOCX_AVAILABLE:
        return

    doc = Document(docx_path)

    # 遍历文档中定义的所有样式
    for style in doc.styles:
        # 寻找名称中包含 "Heading" 或 "标题" 的样式
        if 'Heading' in style.name or '标题' in style.name:
            if hasattr(style, 'font'):
                # 强制设为黑色
                style.font.color.rgb = RGBColor(0, 0, 0)
                # 如果你想让标题也使用宋体/Times New Roman，可以在这里一起设置
                style.font.name = 'Times New Roman'
                style._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

    doc.save(docx_path)

def _html_table_to_markdown(html_table: str) -> str:
    """
    将 HTML 表格转换为 Markdown 管道表格
    支持 rowspan 和 colspan（简化处理：扩展为多个单元格）
    """
    import re
    from html import unescape
    
    # 提取所有行
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html_table, re.DOTALL | re.IGNORECASE)
    if not rows:
        return html_table  # 无法解析，返回原始内容
    
    table_data = []
    for row_html in rows:
        # 提取单元格（th 或 td）
        cells = re.findall(r'<(th|td)[^>]*(?:colspan\s*=\s*["\']?(\d+)["\']?)?[^>]*(?:rowspan\s*=\s*["\']?(\d+)["\']?)?[^>]*>(.*?)</\1>', row_html, re.DOTALL | re.IGNORECASE)
        if not cells:
            # 尝试更宽松的匹配
            cells = re.findall(r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>', row_html, re.DOTALL | re.IGNORECASE)
            cells = [('td', '', '', c) for c in cells]
        
        row_data = []
        for cell in cells:
            if len(cell) == 4:
                tag, colspan, rowspan, content = cell
            else:
                content = cell[0] if cell else ''
                colspan, rowspan = '', ''
            
            # 清理单元格内容
            content = re.sub(r'<[^>]+>', '', content)  # 移除 HTML 标签
            content = unescape(content)  # 解码 HTML 实体
            content = ' '.join(content.split())  # 规范化空白
            content = content.replace('|', '\\|')  # 转义管道符
            
            # 处理 colspan
            col_count = int(colspan) if colspan else 1
            for _ in range(col_count):
                row_data.append(content if _ == 0 else '')
        
        if row_data:
            table_data.append(row_data)
    
    if not table_data:
        return html_table
    
    # 确定最大列数
    max_cols = max(len(row) for row in table_data)
    
    # 补齐每行的列数
    for row in table_data:
        while len(row) < max_cols:
            row.append('')
    
    # 生成 Markdown 表格
    md_lines = []
    for i, row in enumerate(table_data):
        md_lines.append('| ' + ' | '.join(row) + ' |')
        if i == 0:
            # 添加分隔行
            md_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
    
    return '\n' + '\n'.join(md_lines) + '\n'


def _preprocess_html_tables(markdown_content: str) -> str:
    """
    预处理 Markdown 内容，将 HTML 表格转换为 Markdown 表格
    """
    import re
    
    # 匹配 HTML 表格
    table_pattern = r'<table[^>]*>.*?</table>'
    
    def replace_table(match):
        html_table = match.group(0)
        try:
            md_table = _html_table_to_markdown(html_table)
            return md_table
        except Exception as e:
            logger.warning(f"HTML表格转换失败: {e}")
            return html_table
    
    return re.sub(table_pattern, replace_table, markdown_content, flags=re.DOTALL | re.IGNORECASE)


def markdown_to_docx(
    markdown_content: str,
    output_path: str,
    title: Optional[str] = None,
    reference_doc: Optional[str] = None,
    image_base_path: Optional[str] = None,
    extra_args: Optional[list] = None
) -> str:
    """
    将 Markdown 转换为 docx (Word文档)
    使用 Pandoc 作为后端，完美支持表格、公式、代码块等
    
    Args:
        markdown_content: Markdown 文本内容
        output_path: 输出docx文件路径
        title: 文档标题（可选）
        reference_doc: 参考文档路径，用于继承样式（可选）
        image_base_path: 图片基础路径，用于解析相对路径的图片
        extra_args: 传递给 Pandoc 的额外参数
        
    Returns:
        str: 输出文件路径
    """
    if not PYPANDOC_AVAILABLE:
        raise ImportError(
            "pypandoc 未安装，Markdown转docx功能不可用。\n"
            "请运行: pip install pypandoc-binary"
        )
    markdown_content = markdown_content.encode('utf-8', errors='ignore').decode('utf-8')
    # 预处理：将 HTML 表格转换为 Markdown 表格
    processed_content = _preprocess_html_tables(markdown_content)
    
    # 准备 Pandoc 参数
    pandoc_args = [
        '--standalone',
        '--wrap=none',
    ]
    
    # 启用表格和数学公式支持
    # pandoc_args.extend(['-f', 'markdown+pipe_tables+grid_tables+multiline_tables+table_captions+tex_math_dollars'])
    pandoc_args.extend(['-f', 'markdown-raw_tex+pipe_tables+grid_tables+multiline_tables+table_captions+tex_math_dollars'])
    
    # 如果指定了参考文档，使用其样式
    if reference_doc and os.path.exists(reference_doc):
        pandoc_args.extend(['--reference-doc', reference_doc])
    
    # 如果指定了标题
    if title:
        pandoc_args.extend(['-M', f'title={title}'])
    
    # 添加额外参数
    if extra_args:
        pandoc_args.extend(extra_args)
    
    # 设置资源路径（用于解析图片等）
    if image_base_path:
        pandoc_args.extend(['--resource-path', image_base_path])
    
    # 确保输出目录存在
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    try:
        # 直接执行 Markdown 到 docx 转换（保持公式正常工作）
        pypandoc.convert_text(
            processed_content,
            'docx',
            format='markdown',
            outputfile=output_path,
            extra_args=pandoc_args
        )
        
        # 后处理：为表格添加边框
        if PYTHON_DOCX_AVAILABLE:
            try:
                # 1. 添加表格边框
                _add_table_borders(output_path)
                # 2. 设置中西双字体
                _set_fonts(output_path, chinese_font='宋体', latin_font='Times New Roman')
                _fix_styles(output_path)  # 修正颜色
            except Exception as e:
                logger.warning(f"添加表格边框失败: {e}")
        
        logger.info(f"Markdown转Word文档已保存到: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Markdown转docx失败: {e}")
        raise


def markdown_file_to_docx(
    markdown_path: str,
    output_path: Optional[str] = None,
    **kwargs
) -> str:
    """
    将 Markdown 文件转换为 docx
    
    Args:
        markdown_path: Markdown 文件路径
        output_path: 输出docx文件路径，默认与输入文件同目录同名
        **kwargs: 传递给 markdown_to_docx 的其他参数
        
    Returns:
        str: 输出文件路径
    """
    markdown_path = Path(markdown_path)
    
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown文件不存在: {markdown_path}")
    
    # 读取 Markdown 内容
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # 默认输出路径
    if output_path is None:
        output_path = str(markdown_path.with_suffix('.docx'))
    
    # 默认图片基础路径为 Markdown 文件所在目录
    if 'image_base_path' not in kwargs:
        kwargs['image_base_path'] = str(markdown_path.parent)
    
    return markdown_to_docx(markdown_content, output_path, **kwargs)
